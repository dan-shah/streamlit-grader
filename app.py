import streamlit as st
import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold
from pydantic import BaseModel, Field
import PyPDF2
import os
from pathlib import Path
import tempfile
import base64
import json
import re
import io
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
import csv

# Set page config
st.set_page_config(
    page_title="AI Assignment Grader",
    page_icon="📚",
    layout="wide"
)

# Initialize session state
if 'api_key' not in st.session_state:
    st.session_state.api_key = None
if 'rubric_improvements' not in st.session_state:
    st.session_state.rubric_improvements = None
if 'grading_advice' not in st.session_state:
    st.session_state.grading_advice = None
if 'use_analysis_in_grading' not in st.session_state:
    st.session_state.use_analysis_in_grading = False
if 'grading_results' not in st.session_state:
    st.session_state.grading_results = None
if 'submission_text' not in st.session_state:
    st.session_state.submission_text = None
if 'assignment_text' not in st.session_state:
    st.session_state.assignment_text = None
if 'assignment_uploaded_file' not in st.session_state:
    st.session_state.assignment_uploaded_file = None
if 'submission_uploaded_file' not in st.session_state:
    st.session_state.submission_uploaded_file = None
if 'results_pdf' not in st.session_state:
    st.session_state.results_pdf = None
if 'results_dict' not in st.session_state:
    st.session_state.results_dict = None

# Define Pydantic models for structured output
class ImprovementSuggestion(BaseModel):
    area: str = Field(..., description="Area of improvement")
    suggestion: str = Field(..., description="Specific suggestion for improvement")

class PointDeduction(BaseModel):
    area: str = Field(..., description="Area where points were deducted")
    points: int = Field(..., description="Number of points deducted")
    reason: str = Field(..., description="Reason for the deduction")

class ConceptImprovement(BaseModel):
    concept: str = Field(..., description="Concept that needs better understanding")
    suggestion: str = Field(..., description="Specific suggestion to improve understanding")

class GradingFeedback(BaseModel):
    numerical_grade: int = Field(..., description="Numerical grade from 0-100", ge=0, le=100)
    overall_assessment: str = Field(..., description="Overall assessment of the submission")
    strengths: list = Field(..., description="List of strengths in the submission")
    point_deductions: list = Field(..., description="Areas where points were deducted")
    concept_improvements: list = Field(..., description="Suggestions to better grasp concepts")

def extract_text_from_pdf(pdf_file):
    """Extract text from a PDF file."""
    text = ""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
    except Exception as e:
        st.error(f"Error extracting text from PDF: {str(e)}")
        return None
    return text

def analyze_rubric(assignment_rubric_text, api_key):
    """Analyze the rubric/assignment to provide improvement recommendations and grading advice."""
    try:
        # Configure the API
        genai.configure(api_key=api_key)
        
        # Initialize the model with Gemini 2.0 Flash
        model = genai.GenerativeModel('gemini-2.0-flash')
        
        # Create the prompt
        prompt = f"""
        You are an expert educator reviewing a grading rubric and assignment. 
        Please analyze the following assignment and rubric to provide:
        
        1. RUBRIC IMPROVEMENT RECOMMENDATIONS: Analyze the rubric critically and suggest specific improvements 
           that would make it clearer and more effective for consistent grading. Focus on structural improvements, 
           clarity enhancements, and adding specific criteria that may be missing.
        
        2. GRADING ADVICE: Provide specific advice for any AI or human grader on how to interpret and apply 
           this rubric consistently. Highlight key points to look for in submissions, potential pitfalls or 
           misconceptions, and advice for fair evaluation.
        
        Assignment and Rubric:
        {assignment_rubric_text}
        
        Format your response with clear section headers "## RUBRIC IMPROVEMENT RECOMMENDATIONS" and "## GRADING ADVICE".
        Be specific, actionable, and concise in your recommendations.
        """
        
        # Generate response
        response = model.generate_content(prompt)
        response_text = response.text
        
        # Extract the two sections
        improvements_section = ""
        advice_section = ""
        
        # Extract Rubric Improvement Recommendations
        improvements_match = re.search(r'## RUBRIC IMPROVEMENT RECOMMENDATIONS(.*?)(?=## GRADING ADVICE|\Z)', response_text, re.DOTALL)
        if improvements_match:
            improvements_section = improvements_match.group(1).strip()
        
        # Extract Grading Advice
        advice_match = re.search(r'## GRADING ADVICE(.*)', response_text, re.DOTALL)
        if advice_match:
            advice_section = advice_match.group(1).strip()
        
        return {
            "improvements": improvements_section,
            "advice": advice_section,
            "full_response": response_text
        }
    except Exception as e:
        st.error(f"Error analyzing rubric: {str(e)}")
        return None

def grade_assignment(assignment_text, solution_text, submission_text, api_key, include_grading_advice=False, grading_advice=None):
    """Grade the assignment using Gemini with structured output."""
    try:
        # Configure the API
        genai.configure(api_key=api_key)
        
        # Initialize the model with Gemini 2.0 Flash
        generation_config = {
            "temperature": 0.2,
            "top_p": 0.8,
            "top_k": 40,
            "max_output_tokens": 8192,
        }
        
        safety_settings = {
            HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
            HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
            HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
            HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
        }
        
        model = genai.GenerativeModel(
            model_name='gemini-2.0-flash',
            generation_config=generation_config,
            safety_settings=safety_settings
        )
        
        # Create the base prompt
        prompt = f"""
        You are an expert teacher grading an assignment. Please grade the following student submission 
        based on the assignment requirements and provided solution.
        
        Assignment Requirements (including rubric):
        {assignment_text}
        
        Solution:
        {solution_text}
        
        Student Submission:
        {submission_text}
        """
        
        # Include grading advice if requested
        if include_grading_advice and grading_advice:
            prompt += f"""
            
            IMPORTANT GRADING ADVICE:
            {grading_advice}
            """
        
        # Add structured output instructions
        prompt += """
        
        Please provide a detailed evaluation focusing on:
        1. Overall grade with clear justification
        2. Specific strengths shown in the submission
        3. EXPLICIT point deductions - exactly where and why points were lost
        4. Concept-focused improvement suggestions that would help the student better understand the material
        
        IMPORTANT: The total points deducted MUST exactly equal (100 - final_grade). For example, if you assign a grade of 85/100, you must show exactly 15 points of deductions with specific reasons.
        
        Your response should be provided as structured JSON following this schema:
        
        class PointDeduction:
            area: str  # Area where points were deducted
            points: int  # Number of points deducted
            reason: str  # Reason for the deduction
        
        class ConceptImprovement:
            concept: str  # Concept that needs better understanding
            suggestion: str  # Specific suggestion to improve understanding
        
        class GradingFeedback:
            numerical_grade: int  # Numerical grade from 0-100
            overall_assessment: str  # Overall assessment of the submission
            strengths: list[str]  # List of strengths in the submission
            point_deductions: list[PointDeduction]  # Areas where points were deducted
            concept_improvements: list[ConceptImprovement]  # Suggestions to better grasp concepts
        
        Please respond with ONLY a valid JSON object following this schema. Make sure your total point deductions logically explain how you arrived at the final grade and EXACTLY add up to (100 - numerical_grade).
        """
        
        # Generate structured response
        response = model.generate_content(prompt)
        text_response = response.text
        
        # Extract the JSON part from the response
        json_match = re.search(r'```json\s*(.*?)\s*```', text_response, re.DOTALL)
        if json_match:
            json_str = json_match.group(1)
        else:
            # Try to find JSON without code blocks
            json_match = re.search(r'\{.*\}', text_response, re.DOTALL)
            if json_match:
                json_str = json_match.group(0)
            else:
                # If no JSON found, return the raw text
                st.warning("Could not find valid JSON in the response. Using raw text instead.")
                return {"raw_response": text_response}
        
        try:
            # Parse the JSON
            data = json.loads(json_str)
            # Validate with Pydantic
            grading_feedback = GradingFeedback(**data)
            return grading_feedback
        except Exception as e:
            st.error(f"Error parsing structured response: {str(e)}")
            st.warning("Falling back to raw response due to parsing error.")
            return {"raw_response": text_response}
            
    except Exception as e:
        st.error(f"Error grading assignment: {str(e)}")
        return None

def load_sample_files():
    """Load sample files from the data directory."""
    data_dir = Path("data")
    try:
        assignment_file = data_dir / "arima_hw5_assignment_and_rubric.pdf"
        solution_file = data_dir / "arima_hw5_solution_perfect.pdf"
        submission_file = data_dir / "arima_hw5_student_Cplus.pdf"
        
        return {
            "assignment": assignment_file,
            "solution": solution_file,
            "submission": submission_file
        }
    except Exception as e:
        st.error(f"Error loading sample files: {str(e)}")
        return None

def get_file_download_link(file_path, link_text):
    """Generate a download link for a file."""
    with open(file_path, "rb") as f:
        data = f.read()
    b64 = base64.b64encode(data).decode()
    href = f'<a href="data:application/pdf;base64,{b64}" download="{file_path.name}">{link_text}</a>'
    return href

def display_grading_results(results):
    """Display grading results in a structured format."""
    if isinstance(results, GradingFeedback):
        # Calculate total deductions first
        total_deducted = 0
        for deduction in results.point_deductions:
            if isinstance(deduction, dict):
                total_deducted += deduction.get('points', 0)
            else:
                try:
                    total_deducted += deduction.points
                except AttributeError:
                    pass
        
        # Calculate final score based on deductions
        final_score = 100 - total_deducted
        
        # Display structured results
        st.markdown(f"### Grade: {final_score}/100")
        
        st.markdown("#### Overall Assessment")
        st.write(results.overall_assessment)
        
        with st.container():
            st.markdown("#### Strengths")
            for i, strength in enumerate(results.strengths, 1):
                st.markdown(f"**{i}.** {strength}")
        
        with st.container():
            st.markdown("#### Point Deductions")
            for i, deduction in enumerate(results.point_deductions, 1):
                # Handle both object and dictionary formats
                if isinstance(deduction, dict):
                    area = deduction.get('area', f'Area {i}')
                    points = deduction.get('points', 0)
                    reason = deduction.get('reason', 'No reason provided')
                    st.markdown(f"**{i}. {area} (-{points} points)**")
                    st.markdown(f"   {reason}")
                else:
                    try:
                        st.markdown(f"**{i}. {deduction.area} (-{deduction.points} points)**")
                        st.markdown(f"   {deduction.reason}")
                    except AttributeError:
                        st.markdown(f"**{i}. Deduction {i}**")
                        st.markdown(f"   {str(deduction)}")
        
        # Point calculation summary
        with st.container():
            st.markdown("#### Grade Calculation")
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("**Starting Points:**")
                st.markdown("**Total Points Deducted:**")
                st.markdown("**Final Score:**")
            with col2:
                st.markdown("100")
                st.markdown(f"-{total_deducted}")
                st.markdown(f"{final_score}")
        
        with st.container():
            st.markdown("#### Concept Improvement Suggestions")
            for i, improvement in enumerate(results.concept_improvements, 1):
                # Handle both object and dictionary formats
                if isinstance(improvement, dict):
                    concept = improvement.get('concept', f'Concept {i}')
                    suggestion = improvement.get('suggestion', 'No suggestion provided')
                    st.markdown(f"**{i}. {concept}**")
                    st.markdown(f"   {suggestion}")
                else:
                    try:
                        st.markdown(f"**{i}. {improvement.concept}**")
                        st.markdown(f"   {improvement.suggestion}")
                    except AttributeError:
                        st.markdown(f"**{i}. Improvement {i}**")
                        st.markdown(f"   {str(improvement)}")
    elif isinstance(results, dict) and "raw_response" in results:
        # Display raw response if structured parsing failed
        st.markdown("### Grading Results")
        st.markdown(results["raw_response"])
    else:
        # Fallback for any other format
        st.markdown("### Grading Results")
        st.write(results)

def export_to_docx(results):
    """Export grading results to a Word document."""
    if not isinstance(results, (GradingFeedback, dict)):
        return None
    
    doc = Document()
    doc.add_heading('Assignment Grading Results', 0)
    
    # Add student information
    doc.add_heading('Grade Information', level=1)
    if isinstance(results, GradingFeedback):
        doc.add_paragraph(f'Grade: {results.numerical_grade}/100')
        
        # Overall Assessment
        doc.add_heading('Overall Assessment', level=1)
        doc.add_paragraph(results.overall_assessment)
        
        # Strengths
        doc.add_heading('Strengths', level=1)
        for i, strength in enumerate(results.strengths, 1):
            doc.add_paragraph(f"{i}. {strength}", style='List Number')
        
        # Point Deductions
        doc.add_heading('Point Deductions', level=1)
        total_deducted = 0
        for i, deduction in enumerate(results.point_deductions, 1):
            if isinstance(deduction, dict):
                area = deduction.get('area', f'Area {i}')
                points = deduction.get('points', 0)
                reason = deduction.get('reason', 'No reason provided')
                total_deducted += points
                p = doc.add_paragraph(f"{i}. {area} (-{points} points): ", style='List Number')
                p.add_run(f"{reason}")
            else:
                try:
                    total_deducted += deduction.points
                    p = doc.add_paragraph(f"{i}. {deduction.area} (-{deduction.points} points): ", style='List Number')
                    p.add_run(f"{deduction.reason}")
                except AttributeError:
                    doc.add_paragraph(f"{i}. {str(deduction)}", style='List Number')
    else:
        # Raw response
        doc.add_paragraph(results.get("raw_response", "No results available"))
    
    # Save to a BytesIO object
    docx_file = io.BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    
    return docx_file

def generate_results_pdf(results):
    """Generate a PDF of the grading results."""
    if results is None:
        return None
    
    try:
        # If we have results already in dict form, use those
        if hasattr(results, 'dict'):
            results_dict = results.dict()
        elif isinstance(results, dict) and 'numerical_grade' in results:
            results_dict = results
        else:
            # For any other type, convert to string and return a simple PDF
            results_dict = {"raw_response": str(results)}
        
        # Save results_dict to session state for later use
        st.session_state.results_dict = results_dict
        
        # Create a BytesIO object to save the PDF
        pdf_buffer = io.BytesIO()
        
        # Create the PDF document
        doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []
        
        # Add custom styles if they don't already exist
        if 'Heading3' not in styles:
            styles.add(ParagraphStyle(name='Heading3', 
                                    parent=styles['Heading2'], 
                                    fontSize=14,
                                    spaceAfter=6))
        
        # Title
        elements.append(Paragraph("Assignment Grading Results", styles['Title']))
        elements.append(Spacer(1, 12))
        
        if 'numerical_grade' in results_dict:
            # Grade
            elements.append(Paragraph(f"Grade: {results_dict['numerical_grade']}/100", styles['Heading1']))
            elements.append(Spacer(1, 12))
            
            # Overall Assessment
            elements.append(Paragraph("Overall Assessment", styles['Heading2']))
            elements.append(Paragraph(results_dict.get('overall_assessment', 'No assessment available'), styles['Normal']))
            elements.append(Spacer(1, 12))
            
            # Strengths
            elements.append(Paragraph("Strengths", styles['Heading2']))
            for i, strength in enumerate(results_dict.get('strengths', []), 1):
                elements.append(Paragraph(f"{i}. {strength}", styles['Normal']))
            elements.append(Spacer(1, 12))
            
            # Point Deductions
            elements.append(Paragraph("Point Deductions", styles['Heading2']))
            total_deducted = 0
            
            for i, deduction in enumerate(results_dict.get('point_deductions', []), 1):
                if isinstance(deduction, dict):
                    area = deduction.get('area', f'Area {i}')
                    points = deduction.get('points', 0)
                    reason = deduction.get('reason', 'No reason provided')
                    total_deducted += points
                    elements.append(Paragraph(f"<b>{i}. {area} (-{points} points)</b>", styles['Normal']))
                    elements.append(Paragraph(f"{reason}", styles['Normal']))
                else:
                    try:
                        # If it's a Pydantic model, convert to dict
                        if hasattr(deduction, 'dict'):
                            deduction = deduction.dict()
                            area = deduction.get('area', f'Area {i}')
                            points = deduction.get('points', 0)
                            reason = deduction.get('reason', 'No reason provided')
                        else:
                            area = getattr(deduction, 'area', f'Area {i}')
                            points = getattr(deduction, 'points', 0)
                            reason = getattr(deduction, 'reason', 'No reason provided')
                        
                        total_deducted += points
                        elements.append(Paragraph(f"<b>{i}. {area} (-{points} points)</b>", styles['Normal']))
                        elements.append(Paragraph(f"{reason}", styles['Normal']))
                    except Exception as e:
                        elements.append(Paragraph(f"{i}. Unknown deduction: {str(deduction)}", styles['Normal']))
            
            elements.append(Spacer(1, 12))
            
            # Grade Calculation
            elements.append(Paragraph("Grade Calculation", styles['Heading2']))
            data = [
                ["Starting Points:", "100"],
                ["Total Points Deducted:", f"-{total_deducted}"],
                ["Final Score:", f"{results_dict['numerical_grade']}"]
            ]
            t = Table(data, colWidths=[300, 100])
            t.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
            ]))
            elements.append(t)
            
            # Calculation discrepancy warning
            expected_score = 100 - total_deducted
            if expected_score != results_dict['numerical_grade']:
                elements.append(Spacer(1, 6))
                elements.append(Paragraph(f"<i>Note: There appears to be a discrepancy in point calculation. Based on deductions, the expected score would be {expected_score}.</i>", styles['Normal']))
            
            elements.append(Spacer(1, 12))
            
            # Improvement Suggestions
            elements.append(Paragraph("Concept Improvement Suggestions", styles['Heading2']))
            for i, improvement in enumerate(results_dict.get('concept_improvements', []), 1):
                if isinstance(improvement, dict):
                    concept = improvement.get('concept', f'Concept {i}')
                    suggestion = improvement.get('suggestion', 'No suggestion provided')
                    elements.append(Paragraph(f"<b>{i}. {concept}</b>", styles['Normal']))
                    elements.append(Paragraph(f"{suggestion}", styles['Normal']))
                else:
                    try:
                        # If it's a Pydantic model, convert to dict
                        if hasattr(improvement, 'dict'):
                            improvement = improvement.dict()
                            concept = improvement.get('concept', f'Concept {i}')
                            suggestion = improvement.get('suggestion', 'No suggestion provided')
                        else:
                            concept = getattr(improvement, 'concept', f'Concept {i}')
                            suggestion = getattr(improvement, 'suggestion', 'No suggestion provided')
                        
                        elements.append(Paragraph(f"<b>{i}. {concept}</b>", styles['Normal']))
                        elements.append(Paragraph(f"{suggestion}", styles['Normal']))
                    except Exception as e:
                        elements.append(Paragraph(f"{i}. Unknown improvement: {str(improvement)}", styles['Normal']))
        elif "raw_response" in results_dict:
            # Raw response
            elements.append(Paragraph(results_dict.get("raw_response", "No results available"), styles['Normal']))
        else:
            # Fallback for any other format
            elements.append(Paragraph("Grading Results", styles['Heading1']))
            elements.append(Paragraph(str(results_dict), styles['Normal']))
        
        # Build the PDF
        doc.build(elements)
        pdf_buffer.seek(0)
        
        return pdf_buffer
    except Exception as e:
        import traceback
        print(f"Error generating PDF: {str(e)}")
        print(traceback.format_exc())
        return None

def export_to_csv(results):
    """Export grading results to a CSV file."""
    if not isinstance(results, GradingFeedback):
        return None
    
    # Create a BytesIO object to save the CSV
    csv_file = io.StringIO()
    writer = csv.writer(csv_file)
    
    # Write header row
    writer.writerow(["Category", "Details"])
    
    # Grade
    writer.writerow(["Grade", f"{results.numerical_grade}/100"])
    
    # Overall Assessment
    writer.writerow(["Overall Assessment", results.overall_assessment])
    
    # Strengths
    for i, strength in enumerate(results.strengths, 1):
        writer.writerow([f"Strength {i}", strength])
    
    # Point Deductions
    total_deducted = 0
    for i, deduction in enumerate(results.point_deductions, 1):
        if isinstance(deduction, dict):
            area = deduction.get('area', f'Area {i}')
            points = deduction.get('points', 0)
            reason = deduction.get('reason', 'No reason provided')
            total_deducted += points
            writer.writerow([f"Deduction {i}", f"{area} (-{points} points): {reason}"])
        else:
            try:
                total_deducted += deduction.points
                writer.writerow([f"Deduction {i}", f"{deduction.area} (-{deduction.points} points): {deduction.reason}"])
            except AttributeError:
                writer.writerow([f"Deduction {i}", str(deduction)])
    
    # Grade Calculation
    writer.writerow(["Starting Points", "100"])
    writer.writerow(["Total Points Deducted", f"-{total_deducted}"])
    writer.writerow(["Final Score", f"{results.numerical_grade}"])
    
    # Improvement Suggestions
    for i, improvement in enumerate(results.concept_improvements, 1):
        if isinstance(improvement, dict):
            concept = improvement.get('concept', f'Concept {i}')
            suggestion = improvement.get('suggestion', 'No suggestion provided')
            writer.writerow([f"Improvement {i}", f"{concept}: {suggestion}"])
        else:
            try:
                writer.writerow([f"Improvement {i}", f"{improvement.concept}: {improvement.suggestion}"])
            except AttributeError:
                writer.writerow([f"Improvement {i}", str(improvement)])
    
    csv_file.seek(0)
    return csv_file

def display_pdf(pdf_file):
    """Display a PDF file using an iframe."""
    if pdf_file is None:
        return
    
    try:
        # Handle different types of PDF inputs
        if isinstance(pdf_file, io.BytesIO):
            # If it's a BytesIO object, just read it
            pdf_file.seek(0)
            pdf_bytes = pdf_file.read()
        elif isinstance(pdf_file, (str, Path)):
            # If it's a file path (string or Path object)
            with open(pdf_file, "rb") as f:
                pdf_bytes = f.read()
        else:
            # If it's a file-like object (e.g., from st.file_uploader)
            # Reset pointer to the beginning of the file
            pdf_file.seek(0)
            pdf_bytes = pdf_file.read()
        
        # Encode as base64
        base64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')
        
        # Create iframe HTML with improved styling for better scrolling
        pdf_display = f'''
        <div style="display: flex; justify-content: center; width: 100%; height: 500px; overflow: hidden;">
            <iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="100%" 
                    style="border: none; overflow: auto;" type="application/pdf"></iframe>
        </div>
        '''
        st.markdown(pdf_display, unsafe_allow_html=True)
    except Exception as e:
        st.error(f"Error displaying PDF: {str(e)}")
        import traceback
        st.error(traceback.format_exc())

# Main app
st.title("📚 AI Assignment Grader")
st.markdown("""
This application uses Google's Generative AI to automatically grade assignments by comparing student submissions 
against assignment requirements and solution PDFs.
""")

# API Key input
api_key = st.text_input("Enter your Google API Key:", type="password", value=st.session_state.api_key)
if api_key:
    st.session_state.api_key = api_key

# Create tabs for different sections
tabs = st.tabs(["Upload & Grade", "Results View", "Export Results"])

with tabs[0]:  # Upload & Grade tab
    # Sample files section
    st.header("Sample Files")
    st.markdown("""
    You can view and download the sample files used for testing the application:
    """)

    sample_files = load_sample_files()
    if sample_files:
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.subheader("Assignment & Rubric")
            st.markdown(get_file_download_link(sample_files["assignment"], "Download Assignment PDF"), unsafe_allow_html=True)
        
        with col2:
            st.subheader("Solution")
            st.markdown(get_file_download_link(sample_files["solution"], "Download Solution PDF"), unsafe_allow_html=True)
        
        with col3:
            st.subheader("Student Submission")
            st.markdown(get_file_download_link(sample_files["submission"], "Download Submission PDF"), unsafe_allow_html=True)

    # File upload section
    st.header("Upload Files")
    use_sample_files = st.checkbox("Use sample files for testing")

    if use_sample_files:
        if sample_files:
            st.success("Sample files loaded successfully!")
            assignment_file = sample_files["assignment"]
            solution_file = sample_files["solution"]
            submission_file = sample_files["submission"]
            st.session_state.assignment_uploaded_file = assignment_file
            st.session_state.submission_uploaded_file = submission_file
        else:
            st.error("Failed to load sample files. Please upload your own files.")
            use_sample_files = False

    if not use_sample_files:
        assignment_file = st.file_uploader("Upload Assignment PDF (includes rubric)", type="pdf")
        solution_file = st.file_uploader("Upload Solution PDF", type="pdf")
        submission_file = st.file_uploader("Upload Student Submission PDF", type="pdf")
        if assignment_file:
            st.session_state.assignment_uploaded_file = assignment_file
        if submission_file:
            st.session_state.submission_uploaded_file = submission_file

    # Add option to analyze rubric
    if assignment_file and st.button("Analyze Rubric/Assignment"):
        if not st.session_state.api_key:
            st.error("Please enter your Google API Key first!")
        else:
            with st.spinner("Analyzing rubric and assignment..."):
                assignment_text = extract_text_from_pdf(assignment_file)
                if assignment_text:
                    st.session_state.assignment_text = assignment_text
                    analysis_result = analyze_rubric(assignment_text, st.session_state.api_key)
                    if analysis_result:
                        st.session_state.rubric_improvements = analysis_result["improvements"]
                        st.session_state.grading_advice = analysis_result["advice"]
                        st.success("Rubric analysis completed!")
                        
                        # Don't display the analysis here since it will be shown in the expander below
                        # Just set the default value for using analysis in grading
                        st.session_state.use_analysis_in_grading = True
                else:
                    st.error("Failed to extract text from the assignment file.")

    # Show the current analysis if it exists
    if st.session_state.rubric_improvements and st.session_state.grading_advice:
        with st.expander("View Rubric Analysis", expanded=True):
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("### Rubric Improvement Recommendations")
                st.markdown(st.session_state.rubric_improvements)
            
            with col2:
                st.markdown("### Grading Advice")
                st.markdown(st.session_state.grading_advice)
            
            # Allow changing the setting
            st.session_state.use_analysis_in_grading = st.checkbox(
                "Use this grading advice when grading submissions", 
                value=st.session_state.use_analysis_in_grading
            )

    # Grade button
    if st.button("Grade Assignment"):
        if not st.session_state.api_key:
            st.error("Please enter your Google API Key first!")
        elif not use_sample_files and (not assignment_file or not solution_file or not submission_file):
            st.error("Please upload all required PDF files!")
        else:
            with st.spinner("Grading assignment..."):
                # Extract text from PDFs
                assignment_text = extract_text_from_pdf(assignment_file)
                solution_text = extract_text_from_pdf(solution_file)
                submission_text = extract_text_from_pdf(submission_file)
                
                # Store texts in session state for later use
                if assignment_text:
                    st.session_state.assignment_text = assignment_text
                if submission_text:
                    st.session_state.submission_text = submission_text
                
                if all([assignment_text, solution_text, submission_text]):
                    # Grade the assignment
                    result = grade_assignment(
                        assignment_text,
                        solution_text,
                        submission_text,
                        st.session_state.api_key,
                        include_grading_advice=st.session_state.use_analysis_in_grading,
                        grading_advice=st.session_state.grading_advice
                    )
                    
                    if result:
                        # Store results in session state
                        st.session_state.grading_results = result
                        
                        # Convert to dict if it's a Pydantic model
                        if hasattr(result, 'dict'):
                            st.session_state.results_dict = result.dict()
                        elif isinstance(result, dict):
                            st.session_state.results_dict = result
                        
                        # Generate results PDF
                        st.session_state.results_pdf = generate_results_pdf(result)
                        
                        st.success("Grading completed! Check the 'Results View' tab to see the results.")
                        # Switch to the results tab
                        st.query_params["tab"] = "results"
                else:
                    st.error("Failed to extract text from one or more PDF files. Please ensure they are text-based PDFs.")

with tabs[1]:  # Results View tab
    st.header("Grading Results")
    
    if st.session_state.grading_results:
        # Create columns for side-by-side view
        col1, col2 = st.columns([1, 1])
        
        with col1:
            # Show the submission with a toggle for assignment view
            view_type = st.radio(
                "View:",
                ["Student Submission", "Assignment & Rubric"]
            )
            
            if view_type == "Student Submission" and st.session_state.submission_uploaded_file:
                st.subheader("Student Submission")
                display_pdf(st.session_state.submission_uploaded_file)
            elif view_type == "Assignment & Rubric" and st.session_state.assignment_uploaded_file:
                st.subheader("Assignment & Rubric")
                display_pdf(st.session_state.assignment_uploaded_file)
            else:
                st.info("No PDF available to display. Please grade an assignment first.")
        
        with col2:
            # Display grading results as PDF
            st.subheader("Grading Output")
            
            # If we have results in dict form, regenerate the PDF each time
            if st.session_state.results_dict:
                try:
                    # Generate a fresh PDF from the stored dict
                    results_pdf = generate_results_pdf(st.session_state.results_dict)
                    if results_pdf:
                        display_pdf(results_pdf)
                    else:
                        # Fallback to displaying structured content if PDF generation fails
                        st.warning("PDF generation failed. Displaying text version instead:")
                        display_grading_results(st.session_state.grading_results)
                except Exception as e:
                    st.error(f"Error displaying results: {str(e)}")
                    # Fallback to displaying structured content
                    display_grading_results(st.session_state.grading_results)
            elif st.session_state.results_pdf:
                # Use the stored PDF if available
                display_pdf(st.session_state.results_pdf)
            else:
                # Last resort fallback
                display_grading_results(st.session_state.grading_results)
    else:
        st.info("No grading results available. Please grade an assignment first.")
        if st.button("Debug Info"):
            st.write("Session State Debug Info:")
            st.write(f"API Key Set: {bool(st.session_state.api_key)}")
            st.write(f"Grading Results: {bool(st.session_state.grading_results)}")
            st.write(f"Results Dict: {bool(st.session_state.results_dict)}")
            st.write(f"Results PDF: {bool(st.session_state.results_pdf)}")
            st.write(f"Assignment Text: {bool(st.session_state.assignment_text)}")
            st.write(f"Submission Text: {bool(st.session_state.submission_text)}")

with tabs[2]:  # Export Results tab
    st.header("Export Results")
    
    if st.session_state.grading_results:
        export_format = st.selectbox(
            "Select export format:",
            ["Word Document (.docx)", "PDF Document (.pdf)", "CSV File (.csv)"]
        )
        
        if st.button("Export Results"):
            if export_format == "Word Document (.docx)":
                docx_file = export_to_docx(st.session_state.grading_results)
                if docx_file:
                    # Create download button
                    st.download_button(
                        label="Download Word Document",
                        data=docx_file,
                        file_name="grading_results.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            
            elif export_format == "PDF Document (.pdf)":
                pdf_file = generate_results_pdf(st.session_state.grading_results)
                if pdf_file:
                    # Create download button
                    st.download_button(
                        label="Download PDF Document",
                        data=pdf_file,
                        file_name="grading_results.pdf",
                        mime="application/pdf"
                    )
            
            elif export_format == "CSV File (.csv)":
                csv_file = export_to_csv(st.session_state.grading_results)
                if csv_file:
                    # Create download button
                    st.download_button(
                        label="Download CSV File",
                        data=csv_file.getvalue(),
                        file_name="grading_results.csv",
                        mime="text/csv"
                    )
    else:
        st.info("No grading results available for export. Please grade an assignment first.")

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center'>
    <p>Built with ❤️ using Streamlit and Google's Generative AI</p>
    <p>For more information, visit the <a href='https://github.com/dan-shah/streamlit-grader'>GitHub repository</a></p>
</div>
""", unsafe_allow_html=True) 