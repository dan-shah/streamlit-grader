import PyPDF2
import io
import base64
import json
import re
import os
from pathlib import Path
import tempfile
from typing import Optional, Dict, Any, List, BinaryIO
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.shared import Inches
import traceback

# PDF Processing
def extract_text_from_pdf(pdf_file: BinaryIO) -> Optional[str]:
    """Extract text from a PDF file."""
    text = ""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        if len(pdf_reader.pages) == 0:
            raise ValueError("PDF file is empty")
            
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if not page_text:
                raise ValueError("Could not extract text from PDF - it may be scanned or image-based")
            text += page_text + "\n"
        return text
    except Exception as e:
        error_msg = f"Error extracting text from PDF: {str(e)}"
        print(error_msg)
        raise ValueError(error_msg)

# File Operations
def save_temp_file(file_content: bytes, extension: str = ".pdf") -> str:
    """Save bytes content to a temporary file and return the path."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as temp_file:
        temp_file.write(file_content)
        return temp_file.name

def read_file(file_path: str) -> bytes:
    """Read a file and return its contents as bytes."""
    with open(file_path, "rb") as f:
        return f.read()

def encode_file_to_base64(file_content: bytes) -> str:
    """Encode file content to base64 string."""
    return base64.b64encode(file_content).decode('utf-8')

# JSON Parsing
def extract_json_from_text(text: str) -> Optional[Dict[str, Any]]:
    """Extract JSON from text response."""
    # Try to find JSON in code blocks
    json_match = re.search(r'```json\s*(.*?)\s*```', text, re.DOTALL)
    if json_match:
        json_str = json_match.group(1)
    else:
        # Try to find JSON without code blocks
        json_match = re.search(r'\{.*\}', text, re.DOTALL)
        if json_match:
            json_str = json_match.group(0)
        else:
            # If no JSON found, return None
            return None
    
    try:
        # Parse the JSON
        return json.loads(json_str)
    except Exception as e:
        print(f"Error parsing JSON: {str(e)}")
        return None

# Document Generation
def generate_results_pdf(results: Dict[str, Any]) -> Optional[io.BytesIO]:
    """Generate a PDF of the grading results."""
    try:
        # Create a BytesIO object to save the PDF
        pdf_buffer = io.BytesIO()
        
        # Create the PDF document
        doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []
        
        # Add custom styles if they don't already exist
        if 'Heading3' not in styles:
            styles.add(ParagraphStyle(name='Heading3', 
                                    parent=styles['Heading2'], 
                                    fontSize=14,
                                    spaceAfter=6))
        
        # Title
        elements.append(Paragraph("Assignment Grading Results", styles['Title']))
        elements.append(Spacer(1, 12))
        
        if 'numerical_grade' in results:
            # Grade
            elements.append(Paragraph(f"Grade: {results['numerical_grade']}/100", styles['Heading1']))
            elements.append(Spacer(1, 12))
            
            # Overall Assessment
            elements.append(Paragraph("Overall Assessment", styles['Heading2']))
            elements.append(Paragraph(results.get('overall_assessment', 'No assessment available'), styles['Normal']))
            elements.append(Spacer(1, 12))
            
            # Strengths
            elements.append(Paragraph("Strengths", styles['Heading2']))
            for i, strength in enumerate(results.get('strengths', []), 1):
                elements.append(Paragraph(f"{i}. {strength}", styles['Normal']))
            elements.append(Spacer(1, 12))
            
            # Point Deductions
            elements.append(Paragraph("Point Deductions", styles['Heading2']))
            total_deducted = 0
            
            for i, deduction in enumerate(results.get('point_deductions', []), 1):
                if isinstance(deduction, dict):
                    area = deduction.get('area', f'Area {i}')
                    points = deduction.get('points', 0)
                    reason = deduction.get('reason', 'No reason provided')
                    total_deducted += points
                    elements.append(Paragraph(f"<b>{i}. {area} (-{points} points)</b>", styles['Normal']))
                    elements.append(Paragraph(f"{reason}", styles['Normal']))
                else:
                    try:
                        # If it's a Pydantic model, convert to dict
                        if hasattr(deduction, 'dict'):
                            deduction = deduction.dict()
                            area = deduction.get('area', f'Area {i}')
                            points = deduction.get('points', 0)
                            reason = deduction.get('reason', 'No reason provided')
                        else:
                            area = getattr(deduction, 'area', f'Area {i}')
                            points = getattr(deduction, 'points', 0)
                            reason = getattr(deduction, 'reason', 'No reason provided')
                        
                        total_deducted += points
                        elements.append(Paragraph(f"<b>{i}. {area} (-{points} points)</b>", styles['Normal']))
                        elements.append(Paragraph(f"{reason}", styles['Normal']))
                    except Exception as e:
                        elements.append(Paragraph(f"{i}. Unknown deduction: {str(deduction)}", styles['Normal']))
            
            elements.append(Spacer(1, 12))
            
            # Grade Calculation
            elements.append(Paragraph("Grade Calculation", styles['Heading2']))
            data = [
                ["Starting Points:", "100"],
                ["Total Points Deducted:", f"-{total_deducted}"],
                ["Final Score:", f"{results['numerical_grade']}"]
            ]
            t = Table(data, colWidths=[300, 100])
            t.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
            ]))
            elements.append(t)
            
            # Calculation discrepancy warning
            expected_score = 100 - total_deducted
            if expected_score != results['numerical_grade']:
                elements.append(Spacer(1, 6))
                elements.append(Paragraph(f"<i>Note: There appears to be a discrepancy in point calculation. Based on deductions, the expected score would be {expected_score}.</i>", styles['Normal']))
            
            elements.append(Spacer(1, 12))
            
            # Improvement Suggestions
            elements.append(Paragraph("Concept Improvement Suggestions", styles['Heading2']))
            for i, improvement in enumerate(results.get('concept_improvements', []), 1):
                if isinstance(improvement, dict):
                    concept = improvement.get('concept', f'Concept {i}')
                    suggestion = improvement.get('suggestion', 'No suggestion provided')
                    elements.append(Paragraph(f"<b>{i}. {concept}</b>", styles['Normal']))
                    elements.append(Paragraph(f"{suggestion}", styles['Normal']))
                else:
                    try:
                        # If it's a Pydantic model, convert to dict
                        if hasattr(improvement, 'dict'):
                            improvement = improvement.dict()
                            concept = improvement.get('concept', f'Concept {i}')
                            suggestion = improvement.get('suggestion', 'No suggestion provided')
                        else:
                            concept = getattr(improvement, 'concept', f'Concept {i}')
                            suggestion = getattr(improvement, 'suggestion', 'No suggestion provided')
                        
                        elements.append(Paragraph(f"<b>{i}. {concept}</b>", styles['Normal']))
                        elements.append(Paragraph(f"{suggestion}", styles['Normal']))
                    except Exception as e:
                        elements.append(Paragraph(f"{i}. Unknown improvement: {str(improvement)}", styles['Normal']))
        elif "raw_response" in results:
            # Raw response
            elements.append(Paragraph(results.get("raw_response", "No results available"), styles['Normal']))
        else:
            # Fallback for any other format
            elements.append(Paragraph("Grading Results", styles['Heading1']))
            elements.append(Paragraph(str(results), styles['Normal']))
        
        # Build the PDF
        doc.build(elements)
        pdf_buffer.seek(0)
        
        return pdf_buffer
    except Exception as e:
        print(f"Error generating PDF: {str(e)}")
        return None

def export_to_docx(results: Dict[str, Any]) -> Optional[io.BytesIO]:
    """Export grading results to a Word document."""
    try:
        print(f"Starting DOCX generation with results: {results}")  # Debug log
        doc = Document()
        doc.add_heading('Assignment Grading Results', 0)
        
        # Add student information
        doc.add_heading('Grade Information', level=1)
        
        # Convert Pydantic model to dict if needed
        if hasattr(results, 'dict'):
            results = results.dict()
        
        if 'numerical_grade' in results:
            print(f"Processing numerical grade: {results['numerical_grade']}")  # Debug log
            doc.add_paragraph(f'Grade: {results["numerical_grade"]}/100')
            
            # Overall Assessment
            doc.add_heading('Overall Assessment', level=1)
            overall_assessment = results.get('overall_assessment', 'No assessment available')
            print(f"Processing overall assessment: {overall_assessment}")  # Debug log
            doc.add_paragraph(overall_assessment)
            
            # Strengths
            doc.add_heading('Strengths', level=1)
            strengths = results.get('strengths', [])
            print(f"Processing strengths: {strengths}")  # Debug log
            for i, strength in enumerate(strengths, 1):
                doc.add_paragraph(f"{i}. {strength}", style='List Number')
            
            # Point Deductions
            doc.add_heading('Point Deductions', level=1)
            total_deducted = 0
            deductions = results.get('point_deductions', [])
            print(f"Processing deductions: {deductions}")  # Debug log
            for i, deduction in enumerate(deductions, 1):
                try:
                    if isinstance(deduction, dict):
                        area = deduction.get('area', f'Area {i}')
                        points = deduction.get('points', 0)
                        reason = deduction.get('reason', 'No reason provided')
                    else:
                        # If it's a Pydantic model, convert to dict
                        if hasattr(deduction, 'dict'):
                            deduction_dict = deduction.dict()
                            area = deduction_dict.get('area', f'Area {i}')
                            points = deduction_dict.get('points', 0)
                            reason = deduction_dict.get('reason', 'No reason provided')
                        else:
                            area = getattr(deduction, 'area', f'Area {i}')
                            points = getattr(deduction, 'points', 0)
                            reason = getattr(deduction, 'reason', 'No reason provided')
                    
                    total_deducted += points
                    p = doc.add_paragraph(f"{i}. {area} (-{points} points): ", style='List Number')
                    p.add_run(f"{reason}")
                except Exception as e:
                    print(f"Error processing deduction {i}: {str(e)}")  # Debug log
                    doc.add_paragraph(f"{i}. {str(deduction)}", style='List Number')
            
            # Grade Calculation
            doc.add_heading('Grade Calculation', level=1)
            doc.add_paragraph(f'Starting Points: 100')
            doc.add_paragraph(f'Total Points Deducted: -{total_deducted}')
            doc.add_paragraph(f'Final Score: {results["numerical_grade"]}')
            
            # Calculation discrepancy warning
            expected_score = 100 - total_deducted
            if expected_score != results['numerical_grade']:
                doc.add_paragraph(f'Note: There appears to be a discrepancy in point calculation. Based on deductions, the expected score would be {expected_score}.')
            
            # Improvement Suggestions
            doc.add_heading('Concept Improvement Suggestions', level=1)
            improvements = results.get('concept_improvements', [])
            print(f"Processing improvements: {improvements}")  # Debug log
            for i, improvement in enumerate(improvements, 1):
                try:
                    if isinstance(improvement, dict):
                        concept = improvement.get('concept', f'Concept {i}')
                        suggestion = improvement.get('suggestion', 'No suggestion provided')
                    else:
                        # If it's a Pydantic model, convert to dict
                        if hasattr(improvement, 'dict'):
                            improvement_dict = improvement.dict()
                            concept = improvement_dict.get('concept', f'Concept {i}')
                            suggestion = improvement_dict.get('suggestion', 'No suggestion provided')
                        else:
                            concept = getattr(improvement, 'concept', f'Concept {i}')
                            suggestion = getattr(improvement, 'suggestion', 'No suggestion provided')
                    
                    p = doc.add_paragraph(f"{i}. {concept}: ", style='List Number')
                    p.add_run(f"{suggestion}")
                except Exception as e:
                    print(f"Error processing improvement {i}: {str(e)}")  # Debug log
                    doc.add_paragraph(f"{i}. {str(improvement)}", style='List Number')
        else:
            # Raw response
            raw_response = results.get("raw_response", "No results available")
            print(f"Processing raw response: {raw_response}")  # Debug log
            doc.add_paragraph(raw_response)
        
        # Save to a BytesIO object
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        print("DOCX generation completed successfully")  # Debug log
        return docx_file
    except Exception as e:
        print(f"Error generating DOCX: {str(e)}")
        print(f"Full traceback: {traceback.format_exc()}")  # Debug log
        return None 